from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "vibeperform_explore_workshop_angebot.docx"
LOGO = ROOT / "public" / "vibeperform-logo.png"
HERO = ROOT / "public" / "explore-workshop" / "vibeperform-explore-workshop-hero.png"
ROADMAP = ROOT / "public" / "explore-workshop" / "vibeperform-explore-roadmap.png"

PURPLE = "A855F7"
PURPLE_DARK = "7E22CE"
SLATE_950 = "020617"
SLATE_900 = "0F172A"
SLATE_700 = "334155"
SLATE_600 = "475569"
SLATE_500 = "64748B"
SLATE_200 = "E2E8F0"
SLATE_100 = "F1F5F9"
SLATE_50 = "F8FAFC"
WHITE = "FFFFFF"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run(run, *, size=None, bold=None, color=None, font="Arial", italic=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=SLATE_200, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_picture_alt(inline_shape, title, descr):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", descr)


def paragraph_border_bottom(paragraph, color=PURPLE, size="16", space="10"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def set_paragraph_fill(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_para(doc, text="", *, style=None, size=10.5, color=SLATE_700, bold=False, italic=False, before=0, after=6, line=1.18, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    if level == 1:
        set_run(r, size=17, bold=True, color=SLATE_900)
    elif level == 2:
        set_run(r, size=13, bold=True, color=PURPLE_DARK)
    else:
        set_run(r, size=11.5, bold=True, color=SLATE_900)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    set_run(r, size=10.2, color=SLATE_700)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    set_run(r, size=10.2, color=SLATE_700)
    return p


def add_callout(doc, label, body, fill=SLATE_50):
    table = doc.add_table(rows=1, cols=1)
    mark_header_row(table.rows[0])
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(label)
    set_run(run, size=9.5, bold=True, color=PURPLE_DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.18
    run2 = p2.add_run(body)
    set_run(run2, size=10.2, color=SLATE_700)
    add_para(doc, "", after=4)
    return table


def add_matrix(doc, rows, widths, header_fill=SLATE_100):
    table = doc.add_table(rows=1, cols=len(widths))
    mark_header_row(table.rows[0])
    set_table_width(table, widths)
    header = table.rows[0].cells
    for idx, text in enumerate(rows[0]):
        shade_cell(header[idx], header_fill)
        p = header[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run(r, size=9.3, bold=True, color=SLATE_900)
    for data in rows[1:]:
        cells = table.add_row().cells
        for idx, text in enumerate(data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            r = p.add_run(text)
            set_run(r, size=9.2, color=SLATE_700, bold=(idx == 0))
    set_table_width(table, widths)
    add_para(doc, "", after=3)
    return table


def set_document_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.78)
    sec.bottom_margin = Inches(0.78)
    sec.left_margin = Inches(0.86)
    sec.right_margin = Inches(0.86)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(SLATE_700)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Heading 1", 17, SLATE_900, 18, 9),
        ("Heading 2", 13, PURPLE_DARK, 12, 5),
        ("Heading 3", 11.5, SLATE_900, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(10.2)
        style.font.color.rgb = rgb(SLATE_700)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def header_footer(section):
    section.different_first_page_header_footer = True
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    if LOGO.exists():
        run = hp.add_run()
        shape = run.add_picture(str(LOGO), width=Inches(1.25))
        set_picture_alt(shape, "VibePerform logo", "VibePerform wordmark with purple sparkle symbol")
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "VibePerform | Explore Workshop Angebot | contact@vibeperform.com | www.vibeperform.com"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if fp.runs:
        set_run(fp.runs[0], size=8.2, color=SLATE_500)


def add_cover(doc):
    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_logo.add_run()
        shape = run.add_picture(str(LOGO), width=Inches(1.75))
        set_picture_alt(shape, "VibePerform logo", "VibePerform wordmark with purple sparkle symbol")
        p_logo.paragraph_format.space_after = Pt(22)

    kicker = add_para(doc, "ANGEBOT | EXPLORE WORKSHOP", size=9, color=PURPLE_DARK, bold=True, after=8)
    title = add_para(
        doc,
        "Strategischer Explore Workshop",
        size=30,
        color=SLATE_900,
        bold=True,
        after=3,
        line=1.02,
    )
    subtitle = add_para(
        doc,
        "Von der Analyse zur Vision. Von der Überforderung zu absoluter Klarheit.",
        size=14,
        color=SLATE_600,
        after=18,
        line=1.15,
    )
    paragraph_border_bottom(subtitle, color=PURPLE, size="18", space="12")

    meta = doc.add_table(rows=4, cols=2)
    mark_header_row(meta.rows[0])
    set_table_width(meta, [2200, 7160])
    rows = [
        ("Angebot für", "[Kundenname / Organisation]"),
        ("Leistungsbaustein", "Phase 1: Strategic Discovery Sprint"),
        ("Format", "1 Tag vor Ort oder 2 Tage remote"),
        ("Stand", date.today().strftime("%d.%m.%Y")),
    ]
    for idx, (label, value) in enumerate(rows):
        cells = meta.rows[idx].cells
        for c in cells:
            shade_cell(c, SLATE_50 if idx % 2 == 0 else WHITE)
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run(r, size=9.2, bold=True, color=SLATE_500)
        p2 = cells[1].paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        set_run(r2, size=9.8, color=SLATE_900, bold=(idx == 1))
    add_para(doc, "", after=10)

    if HERO.exists():
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = p_img.add_run().add_picture(str(HERO), width=Inches(6.45))
        set_picture_alt(
            shape,
            "Explore Workshop hero image",
            "Workshop table with leaders, process cards, and AI workflow sketches",
        )
        p_img.paragraph_format.space_after = Pt(10)

    add_callout(
        doc,
        "Kurzfassung",
        "Der Explore Workshop schafft eine belastbare Entscheidungsgrundlage: Wir identifizieren business-relevante KI-Hebel, priorisieren Use Cases nach Wirkung und Umsetzbarkeit und dokumentieren die Voraussetzungen für Phase 2.",
        fill="F6F0FF",
    )
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    header_footer(section)


def add_scope(doc):
    add_heading(doc, "1. Ausgangslage und Zielbild", 1)
    add_para(
        doc,
        "Viele Unternehmen starten mit einzelnen KI-Experimenten, ohne ein gemeinsames Lagebild über Ziele, Prozesse, Daten und technische Voraussetzungen zu haben. Dadurch entstehen Demo-Erfolge, aber selten skalierbarer Nutzen.",
    )
    add_para(
        doc,
        "Der Explore Workshop setzt genau hier an. Führung, Fachbereiche, IT/Data und Prozessverantwortliche entwickeln gemeinsam ein priorisiertes Bild der KI-Potenziale und entscheiden, welcher Use Case in der nächsten Phase konkret designt oder pilotiert werden sollte.",
    )
    add_callout(
        doc,
        "Ziel des Angebots",
        "Nach dem Workshop liegt eine priorisierte KI-Roadmap inklusive Use-Case-Briefs, Value-/Effort-Bewertung, ROI-Annahmen und technischen Voraussetzungen vor.",
    )

    add_heading(doc, "2. Leistungsumfang", 1)
    rows = [
        ("Baustein", "Inhalt", "Ergebnis"),
        ("Vorbereitung", "Zielklärung, Teilnehmerkreis, Prozess- und Datenkontext, vorhandene Initiativen.", "Fokussierte Workshop-Hypothesen und Materialliste."),
        ("Workshop", "Moderierte Analyse von Strategie, digitaler Reife, Datenlage, Prozesshebeln und Umsetzungsbild.", "Gemeinsames Lagebild und bewertete Use-Case-Kandidaten."),
        ("Auswertung", "Verdichtung der Ergebnisse, Priorisierung, ROI-Annahmen, Abhängigkeiten und nächste Schritte.", "Management-taugliche Roadmap für Phase 2."),
    ]
    add_matrix(doc, rows, [1700, 4200, 3460])

    add_heading(doc, "3. Workshop-Agenda", 1)
    rows = [
        ("Zeit", "Modul", "Fokus"),
        ("09:00", "Kick-off und Ziele", "Business-Prioritäten, Erfolgskriterien, Entscheidungslogik."),
        ("09:45", "Prozess- und Problem-Inventur", "Manuelle Übergaben, wiederkehrende Abstimmung, Datenbrüche."),
        ("11:15", "KI-Inspiration mit Realitätscheck", "Relevante Beispiele aus Vertrieb, Service und Operations."),
        ("12:30", "Pause", "Mittagspause und informelle Klärungen."),
        ("13:15", "Use-Case-Briefs", "Business Value, Datenbedarf, technische Machbarkeit, Risiken."),
        ("15:15", "Impact-Effort-Priorisierung", "Bewertung nach Wirkung, Aufwand, Abhängigkeit und Akzeptanz."),
        ("16:30", "Roadmap und Management-Entscheidung", "Top-Use-Case, nächste Phase, Verantwortlichkeiten."),
    ]
    add_matrix(doc, rows, [1250, 2920, 5190])

    add_heading(doc, "4. Bewertungsframework", 1)
    for item in [
        "Strategische Ziele: Welche Geschäftsprioritäten soll KI wirklich unterstützen?",
        "Digitale Reife: Welche Systeme, Routinen und Kompetenzen sind belastbar vorhanden?",
        "Datenlage: Wo liegen Silos, Medienbrüche und Qualitätslücken?",
        "Prozesshebel: Welche Workflows sind wiederkehrend, aufwendig und wirtschaftlich relevant?",
        "Umsetzungsbild: Welche Use Cases sind technisch, wirtschaftlich und organisatorisch prioritär?",
    ]:
        add_bullet(doc, item)


def add_deliverables(doc):
    add_heading(doc, "5. Liefergegenstände", 1)
    rows = [
        ("Deliverable", "Beschreibung"),
        ("Priorisierte KI-Roadmap", "Reihenfolge der wichtigsten Use Cases nach Wirkung, Aufwand, Abhängigkeiten und Umsetzungsrisiken."),
        ("3 Use-Case-Briefs", "Kurzprofile mit Zielprozess, Nutzenhypothese, Datenbedarf, Integrationsannahmen und Risiken."),
        ("Value-Heatmap", "Visuelle Entscheidungshilfe für Management und Fachbereiche."),
        ("ROI-Annahmen", "Konkrete Annahmen zu Zeitgewinn, Qualität, Einsparpotenzial und erwartbarer Wirkung."),
        ("Technische Voraussetzungen", "Blick auf Daten, Systeme, Zugriffe, Integrationen, Sicherheit und organisatorische Vorbereitung."),
        ("Next-Step-Empfehlung", "Empfehlung für Design-Sprint, Pilot-Scope oder vorbereitende Data/Integration Readiness."),
    ]
    add_matrix(doc, rows, [2500, 6860])

    if ROADMAP.exists():
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(8)
        shape = p_img.add_run().add_picture(str(ROADMAP), width=Inches(6.2))
        set_picture_alt(
            shape,
            "Explore Workshop roadmap image",
            "Workshop table with impact-effort matrix, roadmap, and data architecture sketches",
        )

    add_heading(doc, "6. Rollen und Mitwirkung", 1)
    rows = [
        ("Rolle", "Empfohlene Beteiligung"),
        ("Geschäftsführung / Bereichsleitung", "Zielbild, Prioritäten, Entscheidungskriterien, Management-Gate."),
        ("Fachbereiche", "Prozesswissen, Pain Points, Beispieldaten, Bewertung der Alltagstauglichkeit."),
        ("IT / Data / Datenschutz", "Systemlandschaft, Datenzugänge, Sicherheitsrahmen, Integrationsannahmen."),
        ("VibePerform", "Moderation, KI-/Prozessmethodik, Use-Case-Strukturierung, Roadmap und Ergebnisdokumentation."),
    ]
    add_matrix(doc, rows, [2500, 6860])

    add_heading(doc, "7. Annahmen und Rahmen", 1)
    for item in [
        "Teilnehmerkreis: idealerweise 6-12 Personen aus Business, IT/Data und relevanten Fachbereichen.",
        "Vorabmaterial: grobe Prozessskizzen, bekannte Engpässe, vorhandene KI-/Automationsideen und Systemübersicht.",
        "Format: 1 Tag vor Ort in und um München oder 2 halbe Tage remote.",
        "Sprache: Deutsch oder Englisch nach Bedarf.",
        "Investition: Festpreis nach finaler Abstimmung von Format, Teilnehmerkreis und Vorbereitungsaufwand.",
    ]:
        add_bullet(doc, item)

    add_callout(
        doc,
        "Nicht enthalten",
        "Technische Implementierung, Datenintegration, produktive Automatisierung und Roll-out sind nicht Bestandteil dieses Angebots. Diese Schritte werden nach dem Explore Workshop als Design-Sprint oder Pilot separat geplant.",
        fill=SLATE_50,
    )


def add_next_steps(doc):
    add_heading(doc, "8. Nächste Schritte", 1)
    for item in [
        "30-minütiges Erstgespräch zur Zielklärung und Bestätigung des passenden Formats.",
        "Abstimmung von Teilnehmerkreis, Workshop-Termin und benötigtem Vorabmaterial.",
        "Finalisierung von Angebot, Zeitplan und organisatorischem Setup.",
        "Durchführung des Explore Workshops und Übergabe der Ergebnisdokumentation.",
    ]:
        add_numbered(doc, item)

    add_heading(doc, "Kontakt", 1)
    add_para(doc, "VibePerform - KI-Beratung für den Mittelstand", size=11, color=SLATE_900, bold=True, after=2)
    add_para(doc, "contact@vibeperform.com | www.vibeperform.com", size=10.5, color=PURPLE_DARK, bold=True, after=12)
    add_callout(
        doc,
        "Ready, steady, deploy.",
        "Von der Idee zur Wirkung - in Wochen, nicht Monaten.",
        fill="F6F0FF",
    )


def main():
    doc = Document()
    set_document_styles(doc)
    header_footer(doc.sections[0])
    add_cover(doc)
    add_scope(doc)
    add_deliverables(doc)
    add_next_steps(doc)

    props = doc.core_properties
    props.title = "VibePerform Explore Workshop Angebot"
    props.subject = "Angebot fuer den strategischen Explore Workshop"
    props.author = "VibePerform"
    props.comments = "Generated offer document for Explore Workshop"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
